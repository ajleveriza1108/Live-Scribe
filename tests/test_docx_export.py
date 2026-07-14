from pathlib import Path

from docx import Document

from src.taglish_transcriber.models import TranscriptSegment
from src.taglish_transcriber.review_engine import ReviewComment
from src.taglish_transcriber.transcript import TranscriptDocument


def test_formatted_docx_contains_final_review_and_live_appendix(tmp_path: Path) -> None:
    transcript = TranscriptDocument()
    transcript.add_live(TranscriptSegment(0.5, 3.0, "Magandang umaga welcome sa meeting"))
    recording = tmp_path / "session.wav"
    enhanced = tmp_path / "session_enhanced.wav"
    recording.write_bytes(b"RIFF")
    enhanced.write_bytes(b"RIFF")
    transcript.set_final(
        [TranscriptSegment(0.5, 3.0, "Magandang umaga. Welcome sa meeting.")],
        [
            ReviewComment(
                timestamp=0.5,
                category="Formatting",
                original="welcome",
                suggestion="Welcome",
                explanation="Possible sentence capitalization.",
            )
        ],
        recording_path=recording,
        enhanced_recording_path=enhanced,
    )

    output = tmp_path / "transcript.docx"
    transcript.save_docx(
        output,
        include_timestamps=True,
        title="Team Meeting Transcript",
        language="Auto — English + Tagalog",
        model="small",
        microphone="Built-in microphone",
    )

    assert output.is_file()
    document = Document(output)
    paragraphs = "\n".join(paragraph.text for paragraph in document.paragraphs)
    tables = "\n".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )
    combined = paragraphs + "\n" + tables
    assert "Team Meeting Transcript" in combined
    assert "Final Reviewed Transcript" in combined
    assert "Magandang umaga. Welcome" in combined
    assert "Grammar, Diction, and Accuracy Comments" in combined
    assert "Original Live Transcript" in combined
    assert "session.wav" in combined
    assert "End of transcript" in combined
    assert "AI-assisted transcription can make mistakes" in combined
