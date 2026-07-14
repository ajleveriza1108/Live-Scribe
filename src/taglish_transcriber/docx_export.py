from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable, Protocol

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


class EntryLike(Protocol):
    start: float
    end: float
    text: str


class ReviewLike(Protocol):
    timestamp: float
    category: str
    original: str
    suggestion: str
    explanation: str
    severity: str


@dataclass(frozen=True, slots=True)
class DocxExportInfo:
    title: str = "Live Transcription"
    language: str = "Auto — English + Tagalog"
    model: str = "small"
    microphone: str = "System default"
    recording_name: str = "Not available"
    enhanced_recording_name: str = "Not created"
    final_accuracy_pass: bool = False
    created_at: datetime | None = None


def _clock(seconds: float) -> str:
    total_seconds = max(0, int(seconds))
    hours, remainder = divmod(total_seconds, 3600)
    minutes, seconds = divmod(remainder, 60)
    return f"{hours:02d}:{minutes:02d}:{seconds:02d}"


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top: int = 90, start: int = 110, bottom: int = 90, end: int = 110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(105, 105, 105)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def _configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.15

    title = document.styles["Title"]
    title.font.name = "Arial"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor(31, 41, 55)

    heading = document.styles["Heading 1"]
    heading.font.name = "Arial"
    heading.font.size = Pt(14)
    heading.font.bold = True
    heading.font.color.rgb = RGBColor(31, 41, 55)
    heading.paragraph_format.space_before = Pt(14)
    heading.paragraph_format.space_after = Pt(7)

    if "Transcript Entry" not in document.styles:
        transcript_style = document.styles.add_style("Transcript Entry", WD_STYLE_TYPE.PARAGRAPH)
    else:
        transcript_style = document.styles["Transcript Entry"]
    transcript_style.font.name = "Arial"
    transcript_style.font.size = Pt(11)
    transcript_style.paragraph_format.space_after = Pt(8)
    transcript_style.paragraph_format.line_spacing = 1.15
    transcript_style.paragraph_format.keep_together = True


def _add_transcript_section(
    document: Document,
    heading: str,
    entries: Iterable[EntryLike],
    *,
    include_timestamps: bool,
) -> None:
    entry_list = list(entries)
    document.add_paragraph(heading, style="Heading 1")
    if not entry_list:
        paragraph = document.add_paragraph("No transcript content was recorded.")
        paragraph.runs[0].italic = True
        paragraph.runs[0].font.color.rgb = RGBColor(100, 100, 100)
        return

    for entry in entry_list:
        paragraph = document.add_paragraph(style="Transcript Entry")
        if include_timestamps:
            timestamp = paragraph.add_run(f"[{_clock(entry.start)}]  ")
            timestamp.bold = True
            timestamp.font.size = Pt(9)
            timestamp.font.color.rgb = RGBColor(91, 99, 112)
        paragraph.add_run(entry.text)


def _add_review_section(document: Document, comments: Iterable[ReviewLike]) -> None:
    comment_list = list(comments)
    document.add_paragraph("Grammar, Diction, and Accuracy Comments", style="Heading 1")
    note = document.add_paragraph(
        "These are review comments only. The app does not silently change what the speaker said."
    )
    note.runs[0].italic = True

    if not comment_list:
        document.add_paragraph("No automatic review comments were generated.")
        return

    table = document.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ("Time", "Category", "Original", "Suggestion", "Comment")
    for index, header in enumerate(headers):
        cell = table.cell(0, index)
        _set_cell_shading(cell, "E9EEF5")
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.size = Pt(9)

    for comment in comment_list:
        row = table.add_row().cells
        values = (
            _clock(comment.timestamp),
            f"{comment.category} ({comment.severity})",
            comment.original,
            comment.suggestion,
            comment.explanation,
        )
        for index, value in enumerate(values):
            row[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            _set_cell_margins(row[index])
            run = row[index].paragraphs[0].add_run(value)
            run.font.size = Pt(8.5)


def save_formatted_docx(
    path: Path,
    *,
    final_entries: Iterable[EntryLike],
    live_entries: Iterable[EntryLike],
    review_comments: Iterable[ReviewLike],
    include_timestamps: bool,
    info: DocxExportInfo,
) -> None:
    final_list = list(final_entries)
    live_list = list(live_entries)
    created_at = info.created_at or datetime.now()
    duration = max((entry.end for entry in final_list or live_list), default=0.0)

    document = Document()
    _configure_styles(document)

    section = document.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    properties = document.core_properties
    properties.title = info.title
    properties.subject = "English, Tagalog, and Taglish transcription"
    properties.author = "Live Scribe"
    properties.keywords = "transcription, English, Tagalog, Filipino, Taglish, WAV"
    properties.comments = "Created locally by Live Scribe."

    header = section.header.paragraphs[0]
    header.text = "LIVE SCRIBE"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.color.rgb = RGBColor(107, 114, 128)
    _add_page_number(section.footer.paragraphs[0])

    title_paragraph = document.add_paragraph(style="Title")
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.add_run(info.title)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle_run = subtitle.add_run("English • Tagalog • Taglish")
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.color.rgb = RGBColor(90, 90, 90)

    table = document.add_table(rows=4, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = (
        "Created", "Language mode", "AI model", "Duration",
        "Microphone", "Final accuracy pass", "Original WAV", "Enhanced WAV",
    )
    values = (
        created_at.strftime("%B %d, %Y at %I:%M %p"),
        info.language,
        info.model,
        _clock(duration),
        info.microphone,
        "Completed" if info.final_accuracy_pass else "Not completed",
        info.recording_name,
        info.enhanced_recording_name,
    )
    for pair_index, (label, value) in enumerate(zip(labels, values)):
        row = (pair_index // 4) * 2
        column = pair_index % 4
        label_cell = table.cell(row, column)
        value_cell = table.cell(row + 1, column)
        _set_cell_shading(label_cell, "E9EEF5")
        _set_cell_margins(label_cell)
        _set_cell_margins(value_cell)
        label_run = label_cell.paragraphs[0].add_run(label)
        label_run.bold = True
        label_run.font.size = Pt(8.5)
        value_run = value_cell.paragraphs[0].add_run(value)
        value_run.font.size = Pt(8.5)

    accuracy_notice = document.add_paragraph()
    accuracy_notice.paragraph_format.space_before = Pt(10)
    accuracy_notice.paragraph_format.space_after = Pt(9)
    notice_label = accuracy_notice.add_run("Accuracy notice: ")
    notice_label.bold = True
    notice_text = accuracy_notice.add_run(
        "AI-assisted transcription can make mistakes. Double-check names, numbers, dates, "
        "amounts, quotations, uncommon words, and other important information against the saved WAV."
    )
    notice_text.italic = True

    _add_transcript_section(
        document,
        "Final Reviewed Transcript" if info.final_accuracy_pass else "Transcript",
        final_list,
        include_timestamps=include_timestamps,
    )
    _add_review_section(document, review_comments)

    if live_list and info.final_accuracy_pass:
        document.add_page_break()
        _add_transcript_section(
            document,
            "Appendix: Original Live Transcript",
            live_list,
            include_timestamps=include_timestamps,
        )

    ending = document.add_paragraph()
    ending.paragraph_format.space_before = Pt(14)
    ending.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ending_run = ending.add_run("End of transcript report")
    ending_run.italic = True
    ending_run.font.size = Pt(9)
    ending_run.font.color.rgb = RGBColor(115, 115, 115)

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))
